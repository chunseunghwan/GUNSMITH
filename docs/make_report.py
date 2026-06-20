"""
GUNSMITH 최종 결과 보고서 생성 스크립트
양식: 최종보고서양식(2026-06-03).hwp 기준
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 페이지 설정 ──────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(2.0)

def set_font(run, name='맑은 고딕', size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, bold=False,
             space_before=0, space_after=6, name='맑은 고딕', color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, name=name, size=size, bold=bold, color=color)
    return p

def set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             valign=WD_ALIGN_VERTICAL.CENTER, name='맑은 고딕'):
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if text:
        run = p.add_run(text)
        set_font(run, name=name, size=size, bold=bold)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 표지
# ═══════════════════════════════════════════════════════════════
add_para('')
add_para('')
add_para('데이터베이스 Term Project 최종 결과 보고서',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True,
         name='맑은 고딕', space_before=60, space_after=30)
add_para('')

# 표지 표
tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# 분반
tbl.cell(0,0).merge(tbl.cell(0,0))
set_cell(tbl.cell(0,0), '분  반', bold=True, size=11)
set_cell_bg(tbl.cell(0,0), 'D9D9D9')
c = tbl.cell(0,1).merge(tbl.cell(0,2))
set_cell(c, '1 분반(GA3007-01)', size=11)

# 작품명
set_cell(tbl.cell(1,0), '작 품 명', bold=True, size=11)
set_cell_bg(tbl.cell(1,0), 'D9D9D9')
c = tbl.cell(1,1).merge(tbl.cell(1,2))
set_cell(c, 'GUNSMITH', size=14, bold=True)
tbl.rows[1].height = Cm(2.2)

# 개발기간
set_cell(tbl.cell(2,0), '개발기간', bold=True, size=11)
set_cell_bg(tbl.cell(2,0), 'D9D9D9')
c = tbl.cell(2,1).merge(tbl.cell(2,2))
set_cell(c, '2026년 5월 20일  ~  2026년 6월 21일', size=11)

# 지도교수
set_cell(tbl.cell(3,0), '지도교수', bold=True, size=11)
set_cell_bg(tbl.cell(3,0), 'D9D9D9')
set_cell(tbl.cell(3,1), '컴퓨터공학과', size=11)
set_cell(tbl.cell(3,2), '오병우', size=11)

# 구분 헤더
set_cell(tbl.cell(4,0), '구  분', bold=True, size=11)
set_cell_bg(tbl.cell(4,0), 'D9D9D9')
set_cell(tbl.cell(4,1), '학년', bold=True, size=11)
set_cell_bg(tbl.cell(4,1), 'D9D9D9')
set_cell(tbl.cell(4,2), '학 번         성 명', bold=True, size=11)
set_cell_bg(tbl.cell(4,2), 'D9D9D9')

# 제출자
set_cell(tbl.cell(5,0), '제출자', size=11)
set_cell(tbl.cell(5,1), '2', size=11)
set_cell(tbl.cell(5,2), '20231161        천승환', size=11)

# 컬럼 너비
from docx.shared import Inches
widths = [Cm(3), Cm(4), Cm(7)]
for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

add_para('')
add_para('')

# 서약문
pledge = doc.add_paragraph()
pledge.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pledge.paragraph_format.space_before = Pt(0)
pledge.paragraph_format.space_after  = Pt(0)
run = pledge.add_run(
    '본인은 데이터베이스 Term Project 최종 결과 보고서를 첨부와 같이 제출합니다. 제출한 보고서는 본인이 직접 개발 및 작성하였으며, 거짓이나 부정이 있다면 F학점을 받고 학칙에 의거하여 처벌받겠습니다 (학적부에 등재).')
set_font(run, size=11)

add_para('')
add_para('2026년   6월   21일', align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)
add_para('')
add_para('제 출 자                    천 승 환', align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
add_para('')
add_para('')
add_para('컴퓨터공학심화프로그램', align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 최종 결과 보고서 (2페이지)
# ═══════════════════════════════════════════════════════════════
add_para('최종 결과 보고서', align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, space_after=2)
add_para('GUNSMITH', align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, space_after=20)

add_para('0.  결과 요약', size=13, bold=True, space_before=10, space_after=6)

# (1) SQL 테이블
add_para('(1) SQL을 사용하여 3개 (Entity 2개 + Relationship 1개) 테이블 이상 생성하고 데이터 삽입',
         size=11, space_after=2)
add_para('  - ○ 중  택 일  →  ○  (구현 완료)',  size=11, space_after=1)
add_para('  - (△라면) 구현한 부분과 구현하지 못한 부분 기술  →  해당 없음', size=11, space_after=1)
add_para('  - 증빙 캡쳐 화면: DuckDB SQL 코드 및 테이블 6개 생성 확인 화면', size=11, space_after=8)

# (2) 3-way JOIN
add_para('(2) 세 개 이상의 테이블 join', size=11, space_after=2)
add_para('  - ○ 중  택 일  →  ○  (구현 완료)', size=11, space_after=1)
add_para('  - (△라면) 구현한 부분과 구현하지 못한 부분 기술  →  해당 없음', size=11, space_after=1)
add_para('  - 증빙 캡쳐 화면: WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT Join 출력 화면', size=11, space_after=8)

# (3) Flet GUI
add_para('(3) Flet으로 GUI 구현', size=11, space_after=2)
add_para('  - ○ 중  택 일  →  ○  (구현 완료)', size=11, space_after=1)
add_para('  - (△라면) 구현한 부분과 구현하지 못한 부분 기술  →  해당 없음', size=11, space_after=1)
add_para('  - 증빙 캡쳐 화면: Flet 애플리케이션 실행 화면', size=11, space_after=8)

# (4) Image
add_para('(4) Image 저장 및 출력', size=11, space_after=2)
add_para('  - ○ 중  택 일  →  ○  (구현 완료)', size=11, space_after=1)
add_para('  - (△라면) 구현한 부분과 구현하지 못한 부분 기술  →  해당 없음', size=11, space_after=1)
add_para('  - 증빙 캡쳐 화면: 1. DuckDB SQL 코드 중 image_data(base64) 저장 부분, 2. image 출력 실행 화면', size=11, space_after=8)

# (5) GitHub
add_para('(5) GitHub Public Repository', size=11, space_after=2)
add_para('  - ○ 중  택 일  →  ○  (구현 완료)', size=11, space_after=1)
add_para('  - GitHub Repository 링크: https://github.com/chunseunghwan/GUNSMITH', size=11, space_after=1)
add_para('  - 증빙 캡쳐 화면: GitHub Repository 화면 캡쳐', size=11, space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 1. 서론
# ═══════════════════════════════════════════════════════════════
add_para('1.  서 론', size=14, bold=True, space_before=6, space_after=8)

intro = (
    'PUBG(PlayerUnknown\'s Battlegrounds)는 배틀로얄 장르의 FPS 게임으로, '
    '다양한 총기와 파츠(부품)를 조합하여 전술을 구성하는 것이 핵심 플레이 요소이다. '
    '총기마다 반동(수직·수평·총구 흔들림 등) 특성이 상이하며, 장착하는 파츠(총구·개머리판·손잡이 등)에 따라 '
    '반동 수치가 달라지기 때문에, 이를 체계적으로 저장·조회·분석할 수 있는 데이터베이스 시스템이 필요하다.\n\n'
    '본 프로젝트 GUNSMITH는 PUBG의 총기 스탯 및 반동 데이터를 DuckDB에 저장하고, '
    'Python + Flet GUI를 통해 총기 조회, 파츠 조회, 커스텀 설정 저장, 반동 결과 비교 등의 기능을 제공하는 데스크톱 애플리케이션이다. '
    '또한 총기 이미지와 파츠 이미지를 base64 인코딩하여 DuckDB TEXT 컬럼에 직접 저장함으로써, '
    '외부 파일 의존 없이 DB 단일 파일만으로 이미지를 포함한 전체 데이터를 이식 가능하도록 설계하였다.\n\n'
    '개발 목표는 PUBG 총기 데이터를 관계형 DB로 구조화하고, 이를 바탕으로 파츠 호환성 조회 및 커스텀 반동 계산 기능을 구현하는 것이다. '
    '특히 3개 이상의 테이블 JOIN(WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT)을 실제 서비스에 활용하고, '
    '이미지 데이터를 DB에 직접 저장하는 방식을 채택하였다.'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run(intro)
set_font(run, size=11)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 2. 작품 개요
# ═══════════════════════════════════════════════════════════════
add_para('2.  작품 개요', size=14, bold=True, space_before=6, space_after=8)
add_para('2.1  전체 구성도', size=12, bold=True, space_after=4)
add_para(
    '사용자는 Flet GUI를 통해 GUNSMITH 애플리케이션에 접근한다. '
    '애플리케이션은 DuckDB로부터 총기/파츠/커스텀/반동 데이터를 조회하고, '
    '결과를 GUI에 렌더링한다. 이미지 데이터(base64)도 DB에서 직접 로드하여 화면에 출력한다.\n\n'
    '[전체 구성도 이미지 첨부]',
    size=11, space_after=10)

add_para('2.2  설계구성요소 및 설계제한 요소', size=12, bold=True, space_after=4)
add_para('<설계구성요소 및 설계제한 요소>', size=10, space_after=3)

# 구성요소 표
t2 = doc.add_table(rows=3, cols=11)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['설계 구성요소', '', '', '', '', '', '설계 제한요소', '', '', '', '']
headers2 = ['목표\n설정', '합성', '분석', '구현/\n제작', '시험/\n평가', '결과\n도출',
            '성능', '규격/\n표준', '경제성', '미학', '안정성/\n내구성', '환경']

# 첫 행 병합
t2.cell(0,0).merge(t2.cell(0,5))
set_cell(t2.cell(0,0), '설계 구성요소', bold=True, size=9)
set_cell_bg(t2.cell(0,0), 'D9D9D9')
t2.cell(0,6).merge(t2.cell(0,10))
set_cell(t2.cell(0,6), '설계 제한요소', bold=True, size=9)
set_cell_bg(t2.cell(0,6), 'D9D9D9')

h2_labels = ['목표\n설정', '합성', '분석', '구현/\n제작', '시험/\n평가', '결과\n도출',
             '성능', '규격/\n표준', '경제성', '미학', '안정성/\n내구성']
for i, lbl in enumerate(h2_labels):
    set_cell(t2.cell(1, i), lbl, bold=True, size=8)
    set_cell_bg(t2.cell(1, i), 'D9D9D9')

checks = ['○', '', '○', '○', '', '', '○', '', '○', '○', '○']
for i, ch in enumerate(checks):
    set_cell(t2.cell(2, i), ch, size=10)

add_para('')
add_para('2.2.1  설계 구성 요소', size=11, bold=True, space_after=4)
add_para('(1) 목표 설정', size=11, bold=True, space_after=2)
add_para('  - PUBG 총기·파츠 데이터를 관계형 DB(DuckDB)로 구조화하여, GUI 기반 조회/커스텀 시뮬레이터를 개발한다.', size=11, space_after=6)

add_para('(2) 분석', size=11, bold=True, space_after=2)
add_para('  - PUBG 총기 반동 데이터 분석: 수직 반동, 수평 반동, 총구 들림, 총구 흔들림, 초탄 반동, 반동 회복 등 6가지 반동 수치 항목 정의', size=11, space_after=1)
add_para('  - 파츠 슬롯 분류: Muzzle, Stock, Grip, Magazine, Scope, Foregrip의 6종 슬롯 체계', size=11, space_after=6)

add_para('(3) 구현/제작', size=11, bold=True, space_after=2)
add_para('  - Python 3.13 + Flet 프레임워크로 GUI 구현', size=11, space_after=1)
add_para('  - DuckDB를 임베디드 DB로 사용, Repository 패턴(추상 인터페이스 + DuckDB 구현체)으로 설계', size=11, space_after=1)
add_para('  - 총기/파츠 이미지를 base64로 인코딩하여 DuckDB TEXT 컬럼에 저장', size=11, space_after=8)

add_para('2.2.2  설계 제한 요소', size=11, bold=True, space_after=4)
add_para('(1) 규격/표준', size=11, bold=True, space_after=2)
add_para('  - SQL 표준 사용, DuckDB 지원 SQL 문법 준수', size=11, space_after=1)
add_para('  - 테이블 3개 이상 사용, JOIN 연산 활용 (WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT)', size=11, space_after=1)
add_para('  - Repository 인터페이스 기반 추상화로 DB 교체 가능성 고려', size=11, space_after=6)

add_para('(2) 미학', size=11, bold=True, space_after=2)
add_para('  - 다크 테마(배경 #0d1117) 기반 직관적 UI 설계', size=11, space_after=1)
add_para('  - 총기·파츠 이미지를 DB에서 직접 로드하여 시각적으로 표시', size=11, space_after=6)

add_para('(3) 안정성', size=11, bold=True, space_after=2)
add_para('  - DB 파일 단일 파일(gunsmith.db)로 관리, 이식성 보장', size=11, space_after=1)
add_para('  - base64 이미지 저장으로 외부 파일 경로 의존성 제거', size=11, space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 3. 설계
# ═══════════════════════════════════════════════════════════════
add_para('3.  설계', size=14, bold=True, space_before=6, space_after=8)
add_para('(설계서와 동일할 수 없음. 구현하면서 변경된 설계 내용을 반영하여 전체 내용 기술)', size=10, color=(255,0,0), space_after=10)

add_para('3.1  ERD (Entity-Relationship Diagram)', size=12, bold=True, space_after=4)
add_para(
    'DuckDB에 생성한 6개 테이블의 ERD를 Crow\'s Foot 표기법으로 작성하였다. '
    'WEAPON과 WEAPON_RECOIL은 1:1 관계, WEAPON과 WEAPON_ATTACHMENT_COMPAT은 1:N 관계, '
    'ATTACHMENT와 WEAPON_ATTACHMENT_COMPAT은 1:N 관계이다. '
    'CUSTOM_CONFIG는 WEAPON을 FK로 참조하며, RECOIL_RESULT는 CUSTOM_CONFIG와 1:1 관계이다.',
    size=11, space_after=4)
add_para('[ERD 이미지 첨부 - docs/erd/GUNSMITH_ERD.png]', size=10, color=(255,0,0), space_after=10)

add_para('3.2  Use Case', size=12, bold=True, space_after=4)

add_para('3.2.1  총기 조회 (Weapon View)', size=11, bold=True, space_after=2)
add_para('  사용자가 총기 목록에서 특정 총기를 선택하면, WEAPON LEFT JOIN WEAPON_RECOIL 쿼리를 실행하여 총기 스탯과 반동 정보를 상세 표시한다. 총기 이미지는 DB의 image_data(base64)를 직접 로드하여 출력한다.', size=11, space_after=6)

add_para('3.2.2  파츠 조회 (Attachment View)', size=11, bold=True, space_after=2)
add_para('  슬롯 타입(Muzzle/Stock/Grip/Magazine/Scope/Foregrip)으로 필터링하여 파츠 목록을 표시한다. 파츠 선택 시 반동 보정값 상세 정보와 파츠 이미지를 표시한다.', size=11, space_after=6)

add_para('3.2.3  파츠 호환 조회 (3-Table JOIN)', size=11, bold=True, space_after=2)
add_para('  특정 총기에 장착 가능한 파츠를 조회할 때 아래 3-way JOIN을 사용한다:', size=11, space_after=2)

# SQL 코드 블록
p_sql = doc.add_paragraph()
p_sql.paragraph_format.left_indent = Cm(1.0)
p_sql.paragraph_format.space_before = Pt(4)
p_sql.paragraph_format.space_after  = Pt(4)
run_sql = p_sql.add_run(
    'SELECT w.weapon_name, a.slot_type, a.attachment_name,\n'
    '       c.control_recoil_vertical, c.control_recoil_horizontal,\n'
    '       c.control_muzzle_rise, c.control_muzzle_shake\n'
    'FROM WEAPON w\n'
    'JOIN WEAPON_ATTACHMENT_COMPAT c ON w.weapon_id = c.weapon_id\n'
    'JOIN ATTACHMENT a               ON c.attachment_id = a.attachment_id\n'
    'WHERE w.weapon_name = \'M416\'\n'
    'ORDER BY a.slot_type, a.attachment_name'
)
set_font(run_sql, name='Courier New', size=9)

add_para('[JOIN 쿼리 출력 캡쳐 화면 첨부]', size=10, color=(255,0,0), space_after=10)

add_para('3.2.4  커스텀 설정 저장 (Custom Config)', size=11, bold=True, space_after=2)
add_para('  사용자가 총기에 파츠를 조합하여 커스텀 설정을 저장하고, 반동 계산 결과(RECOIL_RESULT)를 저장하여 추후 비교 가능하도록 한다.', size=11, space_after=6)

add_para('3.3  UI Design', size=12, bold=True, space_after=4)
add_para('  Flet 프레임워크 기반 멀티 탭 구조로 설계하였다. 총기(Weapon), 파츠(Attachment), 커스텀(Custom) 탭으로 구성되며, 각 탭은 좌측 목록 + 우측 상세 정보 2단 레이아웃을 사용한다.', size=11, space_after=4)
add_para('[UI 화면 구성도 / 스크린샷 첨부]', size=10, color=(255,0,0), space_after=10)

add_para('3.4  Sequence Diagram', size=12, bold=True, space_after=4)
add_para('3.4.1  총기 조회 시퀀스', size=11, bold=True, space_after=2)
add_para('  User → GUI(WeaponView) → WeaponService → WeaponQueryRepository → DuckDB\n'
         '  DuckDB → WeaponQueryRepository → WeaponService → GUI(WeaponView) → User', size=11, space_after=6)

add_para('3.5  Repository Interface 설계', size=12, bold=True, space_after=4)
add_para(
    '추상 인터페이스(IWeaponRepository, IAttachmentRepository 등)와 DuckDB 구현체를 분리하여 '
    'DB 종속성을 최소화하였다. WeaponQueryRepository는 JOIN 쿼리 전용 레포지토리로, '
    '3-way JOIN 등 복합 쿼리를 담당한다.',
    size=11, space_after=6)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 4. 구현
# ═══════════════════════════════════════════════════════════════
add_para('4.  구현', size=14, bold=True, space_before=6, space_after=8)
add_para('개발한 시스템에 대해 Use Case별 자세한 설명 (Use Case별로 실행 캡처 화면과 소스 코드 설명)', size=10, color=(255,0,0), space_after=10)

add_para('4.1  Use Case 1: 총기 목록 조회', size=12, bold=True, space_after=4)
add_para('[실행 화면 캡쳐 첨부]', size=10, color=(255,0,0), space_after=4)
add_para('소스 코드 설명:', size=11, bold=True, space_after=2)
add_para(
    '  WeaponView에서 총기 목록을 로드할 때 WeaponService.get_all_with_grade()를 호출한다. '
    'WeaponQueryRepository는 WEAPON LEFT JOIN WEAPON_RECOIL 쿼리를 실행하여 총기명, 종류, '
    '안정성 점수/등급을 함께 반환한다. GUI에서 안정성 등급(S/A/B/C)에 따라 색상으로 구분 표시한다.',
    size=11, space_after=8)

add_para('4.2  Use Case 2: 총기 상세 + 반동 스탯 조회', size=12, bold=True, space_after=4)
add_para('[실행 화면 캡쳐 첨부]', size=10, color=(255,0,0), space_after=4)
add_para('소스 코드 설명:', size=11, bold=True, space_after=2)
add_para(
    '  총기 선택 시 find_weapon_with_recoil(weapon_id)를 호출하여 WEAPON JOIN WEAPON_RECOIL 결과를 반환한다. '
    'image_data(base64 문자열)를 ft.Image(src=image_data)에 직접 전달하여 총기 이미지를 출력한다.',
    size=11, space_after=8)

add_para('4.3  Use Case 3: 파츠 호환 조회 (3-Table JOIN)', size=12, bold=True, space_after=4)
add_para('[실행 화면 캡쳐 첨부]', size=10, color=(255,0,0), space_after=4)
add_para('소스 코드 설명:', size=11, bold=True, space_after=2)
add_para(
    '  find_compat_parts_by_weapon(weapon_id)는 WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT '
    '3-way JOIN을 실행한다. M416 기준 22개 파츠가 슬롯별로 조회되며, 각 파츠의 반동 보정값을 함께 반환한다.',
    size=11, space_after=8)

add_para('4.4  Use Case 4: 이미지 저장 및 출력', size=12, bold=True, space_after=4)
add_para('[DuckDB image_data 저장 SQL 코드 캡쳐 첨부]', size=10, color=(255,0,0), space_after=2)
add_para('[이미지 출력 실행 화면 캡쳐 첨부]', size=10, color=(255,0,0), space_after=4)
add_para('소스 코드 설명:', size=11, bold=True, space_after=2)
add_para(
    '  db_init.py의 _to_base64(path) 함수가 이미지 파일을 읽어 data:image/webp;base64,... 형식의 '
    'Data URI로 변환한다. 이를 WEAPON.image_data 및 ATTACHMENT.image_data TEXT 컬럼에 INSERT한다. '
    '조회 시 해당 문자열을 ft.Image(src=...)에 그대로 전달하면 네트워크 없이 이미지가 렌더링된다.',
    size=11, space_after=8)

# base64 코드 샘플
p_b64 = doc.add_paragraph()
p_b64.paragraph_format.left_indent = Cm(1.0)
p_b64.paragraph_format.space_before = Pt(4)
p_b64.paragraph_format.space_after  = Pt(8)
run_b64 = p_b64.add_run(
    'def _to_base64(path: str) -> str:\n'
    '    ext = os.path.splitext(path)[1].lower().lstrip(\'.\')\n'
    '    mime = {\'webp\': \'image/webp\', \'png\': \'image/png\', ...}.get(ext, \'image/png\')\n'
    '    with open(path, \'rb\') as f:\n'
    '        b64 = base64.b64encode(f.read()).decode(\'ascii\')\n'
    '    return f\'data:{mime};base64,{b64}\''
)
set_font(run_b64, name='Courier New', size=9)

add_para('4.5  Use Case 5: 커스텀 설정 저장 및 반동 결과 계산', size=12, bold=True, space_after=4)
add_para('[실행 화면 캡쳐 첨부]', size=10, color=(255,0,0), space_after=4)
add_para('소스 코드 설명:', size=11, bold=True, space_after=2)
add_para(
    '  CustomConfigRepository.save_config()로 CUSTOM_CONFIG에 커스텀 설정을 저장한다. '
    'RecoilResultRepository.save_result()로 최종 반동 계산 결과를 RECOIL_RESULT에 저장한다. '
    '커스텀 설정 비교 시 CUSTOM_CONFIG LEFT JOIN RECOIL_RESULT LEFT JOIN WEAPON 쿼리를 활용한다.',
    size=11, space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 5. 테스트 및 검증
# ═══════════════════════════════════════════════════════════════
add_para('5.  테스트 및 검증', size=14, bold=True, space_before=6, space_after=8)

add_para('5.1  DuckDB SQL 테이블 생성 및 데이터 삽입 검증', size=12, bold=True, space_after=4)
add_para(
    '  db_init.py 실행 후 DBeaver 또는 Python 콘솔에서 각 테이블의 데이터를 SELECT하여 '
    '데이터가 정상 삽입되었음을 확인하였다. WEAPON 테이블 8개 총기, ATTACHMENT 테이블 25개 파츠, '
    'WEAPON_ATTACHMENT_COMPAT 테이블 100개 이상의 호환 레코드가 생성되었다.',
    size=11, space_after=8)

add_para('5.2  3-Table JOIN 쿼리 검증', size=12, bold=True, space_after=4)
add_para(
    '  docs/join_demo.py 실행 결과, M416 기준 22개 파츠 레코드가 슬롯별로 정렬 출력되었으며 '
    '각 파츠의 반동 보정값이 정확히 조회됨을 확인하였다.',
    size=11, space_after=8)

add_para('5.3  이미지 저장·출력 검증', size=12, bold=True, space_after=4)
add_para(
    '  M416 총기의 image_data 컬럼에 29,007자의 base64 문자열이 저장되었으며, '
    'Flet GUI에서 ft.Image(src=image_data)를 통해 총기 이미지가 정상 출력됨을 확인하였다.',
    size=11, space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 6. 자체 분석과 평가
# ═══════════════════════════════════════════════════════════════
add_para('6.  자체 분석과 평가', size=14, bold=True, space_before=6, space_after=8)

add_para('  본 프로젝트를 통해 관계형 DB 설계의 정규화 원칙, 외래키 관계, JOIN 연산의 실제 활용 방법을 직접 구현하며 이해할 수 있었다. '
    '특히 3-way JOIN(WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT)을 설계하고 실제 서비스에 '
    '적용하는 과정에서 중간 테이블(다대다 관계 해소용) 설계의 중요성을 체감하였다.',
    size=11, space_after=8)

add_para('설계서에 포함되었지만 구현하지 못한 부분:', size=11, bold=True, space_after=2)
add_para('  - 초기 설계에서는 이미지를 로컬 파일 경로(image_path VARCHAR)로 저장하였으나, '
    '이식성 문제로 base64 인코딩(image_data TEXT) 방식으로 변경하였다. '
    '이로 인해 DB 파일 크기가 증가하는 트레이드오프가 있다.',
    size=11, space_after=8)

add_para('어려웠던 점과 해결 방법:', size=11, bold=True, space_after=2)
add_para('  - DuckDB는 MySQL과 달리 AUTO_INCREMENT 미지원 등 일부 SQL 문법 차이가 있었다. '
    'COALESCE(MAX(id),0)+1 패턴으로 수동 ID 증가 로직을 구현하여 해결하였다.', size=11, space_after=2)
add_para('  - Flet에서 base64 이미지를 ft.Image(src=...)로 직접 렌더링하는 방법을 파악하는 데 시간이 걸렸다. '
    'Flet이 data URI 형식을 직접 지원함을 확인 후 적용하였다.', size=11, space_after=8)

add_para('본인 작성 vs AI 활용 부분:', size=11, bold=True, space_after=2)
add_para('  - 전체 DB 설계(테이블 구조, 정규화, 관계 설정), Repository 패턴 구현, GUI 레이아웃 설계는 본인이 직접 작성하였다.', size=11, space_after=2)
add_para('  - Claude AI를 활용하여 base64 인코딩 DB 저장 방식 전환, ERD 다이어그램 생성 코드(matplotlib), 보고서 작성을 보조받았다.', size=11, space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════
# ■ 7. 부록
# ═══════════════════════════════════════════════════════════════
add_para('7.  부록', size=14, bold=True, space_before=6, space_after=8)
add_para('7.1  개발 일지', size=12, bold=True, space_after=4)

# 개발 일지 표
log_data = [
    ('2026.05.20', 'DB 설계 시작 - 요구사항 분석, 테이블 설계 (WEAPON, WEAPON_RECOIL, ATTACHMENT 등 6개 테이블)'),
    ('2026.05.22', 'MySQL 기반 DDL 작성, 정규화 검토'),
    ('2026.05.27', 'DuckDB로 마이그레이션, db_init.py 작성, 샘플 데이터 삽입'),
    ('2026.05.30', 'Repository 패턴 설계 - 추상 인터페이스 및 DuckDB 구현체 작성'),
    ('2026.06.03', 'Flet GUI 기본 구조 구현 - WeaponView, AttachmentView'),
    ('2026.06.07', '3-way JOIN 쿼리 구현 및 검증 (WEAPON ⋈ WEAPON_ATTACHMENT_COMPAT ⋈ ATTACHMENT)'),
    ('2026.06.10', 'CustomConfigView, RecoilResultView 구현'),
    ('2026.06.14', '이미지 저장 방식 변경: image_path → image_data(base64 TEXT)'),
    ('2026.06.18', '전체 테스트 및 버그 수정'),
    ('2026.06.20', 'ERD 다이어그램 생성, 최종 보고서 작성, GitHub push'),
    ('2026.06.21', '최종 제출'),
]

tbl_log = doc.add_table(rows=len(log_data)+1, cols=2)
tbl_log.style = 'Table Grid'
tbl_log.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell(tbl_log.cell(0,0), '날짜', bold=True, size=10)
set_cell_bg(tbl_log.cell(0,0), 'D9D9D9')
set_cell(tbl_log.cell(0,1), '작업 내용', bold=True, size=10)
set_cell_bg(tbl_log.cell(0,1), 'D9D9D9')
for i, (date, content) in enumerate(log_data, 1):
    set_cell(tbl_log.cell(i,0), date, size=10)
    set_cell(tbl_log.cell(i,1), content, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

# 컬럼 너비
for row in tbl_log.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(11)

add_para('')

# ═══════════════════════════════════════════════════════════════
# ■ 참고문헌
# ═══════════════════════════════════════════════════════════════
add_para('참 고 문 헌', size=13, bold=True, space_before=12, space_after=6)
refs = [
    '[1] PUBG 공식 위키, https://pubg.fandom.com',
    '[2] DuckDB 공식 문서, https://duckdb.org/docs',
    '[3] Flet 공식 문서, https://flet.dev/docs',
    '[4] Python 공식 문서, https://docs.python.org/3',
    '[5] GitHub Repository, https://github.com/chunseunghwan/GUNSMITH',
]
for ref in refs:
    add_para(ref, size=10, space_after=2)

# 저장
out = r'C:\Users\user\DB_2026\GUNSMITH\docs\GUNSMITH_최종보고서.docx'
doc.save(out)
print('saved:', out)
